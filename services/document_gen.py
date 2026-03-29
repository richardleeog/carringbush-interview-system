"""
Document generation service for creating professional CVs, cover letters,
and meeting summaries using python-docx and optional Claude API enhancement.

Creates clean, professional documents suitable for entry-level job seekers
in Melbourne across hospitality, aged care, cleaning, and hotel sectors.
"""

import logging
import json
import os
from datetime import datetime
from typing import Optional, Dict, Any
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsmap
    # Register 'w' namespace prefix if not already present
    if 'w' not in nsmap:
        nsmap['w'] = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning('python-docx not installed. Document generation will be limited.')

logger = logging.getLogger(__name__)

# Design colours — guarded for when python-docx is not installed
try:
    COLOUR_DARK_BLUE = RGBColor(0x1B, 0x4F, 0x72)  # #1B4F72
    COLOUR_BLUE = RGBColor(0x2E, 0x86, 0xC1)      # #2E86C1
except NameError:
    COLOUR_DARK_BLUE = None
    COLOUR_BLUE = None
FONT_NAME = 'Calibri'


class DocumentGenerator:
    """Generates professional documents for job seekers."""

    def __init__(self, config=None):
        """
        Initialise the document generator.

        Args:
            config: Configuration object with STUDENT_FILES_DIR and optional
                   ANTHROPIC_API_KEY for content enhancement.
        """
        self.config = config
        if isinstance(config, dict):
            self.student_files_dir = config.get('STUDENT_FILES_DIR', '/tmp/carringbush/students')
        else:
            self.student_files_dir = getattr(
                config, 'STUDENT_FILES_DIR', '/tmp/carringbush/students'
            ) if config else '/tmp/carringbush/students'

        self.api_key = getattr(config, 'ANTHROPIC_API_KEY', None) if config else None
        self._claude_client = None

        if not DOCX_AVAILABLE:
            logger.warning(
                'python-docx not available. Install with: pip install python-docx'
            )

    def _get_claude_client(self):
        """Lazily initialise Claude client if API key is available."""
        if self._claude_client is None and self.api_key:
            try:
                from anthropic import Anthropic
                self._claude_client = Anthropic(api_key=self.api_key)
            except ImportError:
                logger.warning('Anthropic library not installed')
            except Exception as e:
                logger.warning(f'Failed to initialise Claude client: {e}')

        return self._claude_client

    def _polish_content(self, content: str, content_type: str) -> str:
        """
        Use Claude to polish and enhance content if API is available.

        Args:
            content: Raw content to polish
            content_type: Type of content (e.g., 'cv_summary', 'work_experience')

        Returns:
            Polished content, or original content if Claude unavailable
        """
        client = self._get_claude_client()
        if not client:
            return content

        try:
            prompts = {
                'cv_summary': (
                    'You are a professional CV writer. Rewrite this professional summary '
                    'to be compelling, concise, and suitable for entry-level job seekers. '
                    'Keep it under 3 sentences. Do not include any introductory text, '
                    'just the polished summary.\n\n'
                ),
                'work_experience': (
                    'You are a professional CV writer. Rewrite this work experience description '
                    'to use strong action verbs and emphasise achievements. Keep it concise. '
                    'Return only the polished description, no introduction.\n\n'
                ),
                'skills': (
                    'You are a professional CV writer. Improve this list of skills to be more '
                    'professional and impactful. Use standard terminology for job seekers. '
                    'Return only the improved list, no introduction.\n\n'
                )
            }

            prompt = prompts.get(content_type, 'Improve this text: ')
            response = client.messages.create(
                model='claude-3-5-sonnet-20241022',
                max_tokens=500,
                messages=[
                    {'role': 'user', 'content': prompt + content}
                ]
            )
            polished = response.content[0].text.strip()
            logger.info(f'Content polished via Claude: {content_type}')
            return polished
        except Exception as e:
            logger.warning(f'Claude polishing failed for {content_type}: {e}')
            return content

    def _create_document_header(self, doc: Document, student_data: Dict[str, Any]):
        """Create the professional header section of a document."""
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(
            f"{student_data.get('first_name', '')} {student_data.get('surname', '')}"
        )
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.color.rgb = COLOUR_DARK_BLUE
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_parts = []
        if student_data.get('phone'):
            contact_parts.append(student_data['phone'])
        if student_data.get('email'):
            contact_parts.append(student_data['email'])
        if student_data.get('location'):
            contact_parts.append(student_data['location'])

        if contact_parts:
            contact_para = doc.add_paragraph(' • '.join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in contact_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = COLOUR_BLUE

        doc.add_paragraph()  # Spacing

    def _add_section_header(self, doc: Document, heading: str):
        """Add a styled section header."""
        para = doc.add_paragraph()
        run = para.add_run(heading)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOUR_DARK_BLUE

        # Add underline
        pPr = para._element.get_or_add_pPr()
        pBdr = pPr.find('.//w:pBdr')
        if pBdr is None:
            from docx.oxml import OxmlElement
            pBdr = OxmlElement('w:pBdr')
            pPr.append(pBdr)
            bottom = OxmlElement('w:bottom')
            bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
            bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '24')
            bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '1')
            bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '2E86C1')
            pBdr.append(bottom)

    def generate_cv(
        self,
        student_data: Dict[str, Any],
        session_data: Dict[str, Any],
        language: str = 'en'
    ) -> Optional[str]:
        """
        Generate a professional CV document.

        Args:
            student_data: Student profile information
            session_data: Interview session data with cv_content, work_history, skills, etc.
            language: Output language code (e.g., 'en', 'es', 'fr')

        Returns:
            Path to generated CV file, or None if generation failed
        """
        if not DOCX_AVAILABLE:
            logger.error('python-docx not available')
            return None

        try:
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = FONT_NAME
            style.font.size = Pt(11)

            # Header
            self._create_document_header(doc, student_data)

            # Professional Summary
            if session_data.get('cv_content', {}).get('professional_summary'):
                summary = session_data['cv_content']['professional_summary']
                summary = self._polish_content(summary, 'cv_summary')

                self._add_section_header(doc, 'Professional Summary')
                para = doc.add_paragraph(summary)
                para.paragraph_format.space_after = Pt(6)

            # Key Skills (2-column layout)
            skills = session_data.get('cv_content', {}).get('key_skills', [])
            if skills:
                skills_text = ', '.join(skills) if isinstance(skills, list) else skills
                skills_text = self._polish_content(skills_text, 'skills')

                self._add_section_header(doc, 'Key Skills')
                para = doc.add_paragraph(skills_text)
                para.paragraph_format.space_after = Pt(6)

            # Work Experience
            work_history = session_data.get('work_history', [])
            if work_history:
                self._add_section_header(doc, 'Work Experience')

                for job in work_history:
                    # Job title and company
                    job_para = doc.add_paragraph()
                    title_run = job_para.add_run(
                        f"{job.get('title', '')} at {job.get('company', '')}"
                    )
                    title_run.font.bold = True
                    title_run.font.size = Pt(11)
                    title_run.font.color.rgb = COLOUR_BLUE

                    # Dates
                    dates_text = f"{job.get('start_date', '')} - {job.get('end_date', '')}"
                    dates_para = doc.add_paragraph(dates_text)
                    dates_para.paragraph_format.space_after = Pt(3)

                    # Description
                    description = job.get('description', '')
                    if description:
                        description = self._polish_content(description, 'work_experience')
                        desc_para = doc.add_paragraph(description, style='List Bullet')
                        desc_para.paragraph_format.space_after = Pt(6)

            # Education
            education = session_data.get('education', [])
            if education:
                self._add_section_header(doc, 'Education')

                for edu in education:
                    edu_para = doc.add_paragraph()
                    qual_run = edu_para.add_run(edu.get('qualification', ''))
                    qual_run.font.bold = True
                    qual_run.font.color.rgb = COLOUR_BLUE

                    school_para = doc.add_paragraph(edu.get('institution', ''))
                    school_para.paragraph_format.space_after = Pt(3)

                    if edu.get('year'):
                        year_para = doc.add_paragraph(f"Year: {edu.get('year')}")
                        year_para.paragraph_format.space_after = Pt(6)

            # Languages
            languages = session_data.get('languages', [])
            if languages:
                self._add_section_header(doc, 'Languages')
                for lang in languages:
                    para = doc.add_paragraph(
                        f"{lang.get('language', '')}: {lang.get('proficiency', '')}",
                        style='List Bullet'
                    )
                    para.paragraph_format.space_after = Pt(3)

            # Additional Information
            additional_info = session_data.get('cv_content', {}).get('additional_info')
            if additional_info:
                self._add_section_header(doc, 'Additional Information')
                para = doc.add_paragraph(additional_info)
                para.paragraph_format.space_after = Pt(6)

            # References
            references = session_data.get('references', [])
            if references:
                self._add_section_header(doc, 'References')
                para = doc.add_paragraph('Available upon request')
                para.paragraph_format.space_after = Pt(6)

            # Save document
            output_dir = self._get_session_output_dir(student_data, session_data)
            output_path = os.path.join(output_dir, f'CV_{language.upper()}.docx')

            doc.save(output_path)
            logger.info(f'CV generated: {output_path}')
            return output_path

        except Exception as e:
            logger.error(f'CV generation failed: {e}')
            return None

    def generate_cover_letter(
        self,
        student_data: Dict[str, Any],
        session_data: Dict[str, Any],
        job_title: str = '',
        employer: str = '',
        language: str = 'en'
    ) -> Optional[str]:
        """
        Generate a cover letter document.

        Args:
            student_data: Student profile information
            session_data: Interview session data
            job_title: Target job position
            employer: Target employer name
            language: Output language code

        Returns:
            Path to generated cover letter file, or None if generation failed
        """
        if not DOCX_AVAILABLE:
            logger.error('python-docx not available')
            return None

        try:
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = FONT_NAME
            style.font.size = Pt(11)

            # Date
            date_para = doc.add_paragraph(datetime.now().strftime('%d %B %Y'))
            doc.add_paragraph()

            # Recipient
            recipient_para = doc.add_paragraph(f'{employer}\n[Address to be added]')
            doc.add_paragraph()

            # Greeting
            greeting = 'Dear Hiring Manager,' if not job_title else f'Dear {employer} Team,'
            doc.add_paragraph(greeting)

            # Opening paragraph
            opening = (
                f'I am writing to express my interest in the {job_title} position at {employer}. '
                'I am an enthusiastic and reliable professional seeking an opportunity to contribute '
                'to your team and grow professionally.'
            )
            doc.add_paragraph(opening)

            # Skills and experience
            skills = session_data.get('cv_content', {}).get('key_skills', [])
            if skills:
                skills_text = ', '.join(skills) if isinstance(skills, list) else skills
                experience_para = (
                    f'I bring the following skills and attributes: {skills_text}. '
                    'I am committed to delivering quality work and maintaining high standards '
                    'in every role I undertake.'
                )
                doc.add_paragraph(experience_para)

            # Closing
            closing = (
                'I would welcome the opportunity to discuss how I can contribute to your organisation. '
                'Thank you for considering my application. I look forward to hearing from you.'
            )
            doc.add_paragraph(closing)

            # Sign-off
            doc.add_paragraph()
            doc.add_paragraph(f"Yours sincerely,\n\n{student_data.get('first_name', '')} {student_data.get('surname', '')}")

            # Save document
            output_dir = self._get_session_output_dir(student_data, session_data)
            output_path = os.path.join(output_dir, f'CoverLetter_{language.upper()}.docx')

            doc.save(output_path)
            logger.info(f'Cover letter generated: {output_path}')
            return output_path

        except Exception as e:
            logger.error(f'Cover letter generation failed: {e}')
            return None

    def generate_meeting_summary_internal(
        self,
        student_data: Dict[str, Any],
        session_data: Dict[str, Any]
    ) -> Optional[str]:
        """
        Generate an internal meeting summary for staff records.

        Args:
            student_data: Student profile information
            session_data: Interview session data

        Returns:
            Path to generated summary file, or None if generation failed
        """
        if not DOCX_AVAILABLE:
            logger.error('python-docx not available')
            return None

        try:
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = FONT_NAME

            # Header
            title = doc.add_paragraph(f"Meeting Summary: {student_data.get('first_name', '')} {student_data.get('surname', '')}")
            title_run = title.runs[0]
            title_run.font.bold = True
            title_run.font.size = Pt(14)
            title_run.font.color.rgb = COLOUR_DARK_BLUE

            doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
            doc.add_paragraph(f"Session ID: {session_data.get('session_id', 'N/A')}")
            doc.add_paragraph()

            # Student Information
            self._add_section_header(doc, 'Student Information')
            doc.add_paragraph(f"Name: {student_data.get('first_name', '')} {student_data.get('surname', '')}")
            doc.add_paragraph(f"Email: {student_data.get('email', 'N/A')}")
            doc.add_paragraph(f"Phone: {student_data.get('phone', 'N/A')}")
            doc.add_paragraph(f"Location: {student_data.get('location', 'N/A')}")
            doc.add_paragraph()

            # Interview Content Summary
            self._add_section_header(doc, 'Interview Content')
            if session_data.get('cv_content', {}).get('professional_summary'):
                doc.add_paragraph('Professional Summary:')
                doc.add_paragraph(
                    session_data['cv_content']['professional_summary'],
                    style='List Bullet'
                )

            if session_data.get('cv_content', {}).get('key_skills'):
                doc.add_paragraph('Key Skills:')
                skills = session_data['cv_content']['key_skills']
                for skill in (skills if isinstance(skills, list) else [skills]):
                    doc.add_paragraph(skill, style='List Bullet')

            if session_data.get('work_history'):
                doc.add_paragraph('Work History:')
                for job in session_data['work_history']:
                    doc.add_paragraph(
                        f"{job.get('title', '')} at {job.get('company', '')} ({job.get('start_date', '')} - {job.get('end_date', '')})",
                        style='List Bullet'
                    )

            doc.add_paragraph()

            # Notes for staff
            self._add_section_header(doc, 'Staff Notes')
            doc.add_paragraph(session_data.get('notes', 'No notes recorded'))

            # Save document
            output_dir = self._get_session_output_dir(student_data, session_data)
            output_path = os.path.join(output_dir, 'MeetingSummary_Internal.docx')

            doc.save(output_path)
            logger.info(f'Internal meeting summary generated: {output_path}')
            return output_path

        except Exception as e:
            logger.error(f'Internal meeting summary generation failed: {e}')
            return None

    def generate_meeting_summary_student(
        self,
        student_data: Dict[str, Any],
        session_data: Dict[str, Any],
        language: str = 'en'
    ) -> Optional[str]:
        """
        Generate a student-friendly meeting summary.

        Args:
            student_data: Student profile information
            session_data: Interview session data
            language: Output language code

        Returns:
            Path to generated summary file, or None if generation failed
        """
        if not DOCX_AVAILABLE:
            logger.error('python-docx not available')
            return None

        try:
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = FONT_NAME

            # Header
            title = doc.add_paragraph(f"Your Interview Summary - {datetime.now().strftime('%d %B %Y')}")
            title_run = title.runs[0]
            title_run.font.bold = True
            title_run.font.size = Pt(14)
            title_run.font.color.rgb = COLOUR_DARK_BLUE
            doc.add_paragraph()

            # What we discussed
            self._add_section_header(doc, 'What We Discussed')
            doc.add_paragraph(
                'During our interview, we captured the following information for your CV:'
            )
            doc.add_paragraph()

            # Professional summary
            if session_data.get('cv_content', {}).get('professional_summary'):
                doc.add_paragraph('Your Professional Summary:')
                summary = session_data['cv_content']['professional_summary']
                doc.add_paragraph(summary, style='List Bullet')
                doc.add_paragraph()

            # Skills
            if session_data.get('cv_content', {}).get('key_skills'):
                doc.add_paragraph('Your Key Skills:')
                skills = session_data['cv_content']['key_skills']
                for skill in (skills if isinstance(skills, list) else [skills]):
                    doc.add_paragraph(skill, style='List Bullet')
                doc.add_paragraph()

            # Next steps
            self._add_section_header(doc, 'Next Steps')
            doc.add_paragraph(
                'Your CV has been created and is ready for use. '
                'You can download it and use it to apply for jobs. '
                'If you would like any changes, please let us know.'
            )

            # Save document
            output_dir = self._get_session_output_dir(student_data, session_data)
            output_path = os.path.join(output_dir, f'MeetingSummary_Student_{language.upper()}.docx')

            doc.save(output_path)
            logger.info(f'Student meeting summary generated: {output_path}')
            return output_path

        except Exception as e:
            logger.error(f'Student meeting summary generation failed: {e}')
            return None

    def generate_action_items(
        self,
        session_data: Dict[str, Any],
        language: str = 'en'
    ) -> Optional[str]:
        """
        Generate an action items document.

        Args:
            session_data: Interview session data containing action items
            language: Output language code

        Returns:
            Path to generated action items file, or None if generation failed
        """
        if not DOCX_AVAILABLE:
            logger.error('python-docx not available')
            return None

        try:
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = FONT_NAME

            # Header
            title = doc.add_paragraph('Action Items')
            title_run = title.runs[0]
            title_run.font.bold = True
            title_run.font.size = Pt(14)
            title_run.font.color.rgb = COLOUR_DARK_BLUE

            doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y at %H:%M')}")
            doc.add_paragraph()

            # Action items
            action_items = session_data.get('action_items', [])
            if action_items:
                for idx, item in enumerate(action_items, 1):
                    item_text = item if isinstance(item, str) else item.get('description', '')
                    due_date = item.get('due_date', '') if isinstance(item, dict) else ''

                    para = doc.add_paragraph(item_text, style='List Number')
                    if due_date:
                        due_para = doc.add_paragraph(f'Due: {due_date}')
                        due_para.paragraph_format.left_indent = Inches(0.5)
                    doc.add_paragraph()
            else:
                doc.add_paragraph('No action items recorded.')

            # Save document
            output_dir = session_data.get('output_dir', '/tmp/carringbush')
            Path(output_dir).mkdir(parents=True, exist_ok=True)

            output_path = os.path.join(output_dir, f'ActionItems_{language.upper()}.docx')

            doc.save(output_path)
            logger.info(f'Action items generated: {output_path}')
            return output_path

        except Exception as e:
            logger.error(f'Action items generation failed: {e}')
            return None

    def _get_session_output_dir(
        self,
        student_data: Dict[str, Any],
        session_data: Dict[str, Any]
    ) -> str:
        """Get or create the session output directory."""
        surname = student_data.get('surname', 'Unknown').replace(' ', '_')
        first_name = student_data.get('first_name', 'Unknown').replace(' ', '_')
        student_id = student_data.get('id', 'unknown')

        student_dir = os.path.join(
            self.student_files_dir,
            f'{surname}_{first_name}_{student_id}'
        )

        session_date = session_data.get('date', datetime.now().strftime('%Y-%m-%d'))
        session_num = session_data.get('session_number', 1)
        session_dir = os.path.join(student_dir, f'{session_date}_Session_{session_num:02d}')

        Path(session_dir).mkdir(parents=True, exist_ok=True)
        return session_dir
